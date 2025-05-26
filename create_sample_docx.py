#!/usr/bin/env python3
"""
Script to create a sample DOCX file for testing docx-plainify.
"""

from docx import Document
from docx.shared import Inches

def create_sample_docx():
    """Create a sample DOCX file with various elements."""
    
    # Create a new document
    doc = Document()
    
    # Add a title
    title = doc.add_heading('项目计划文档', 0)
    
    # Add a paragraph
    doc.add_paragraph('这是一个示例文档，用于测试 docx-plainify 工具的转换功能。')
    
    # Add a heading
    doc.add_heading('项目概述', level=1)
    
    # Add another paragraph
    doc.add_paragraph('本项目旨在开发一个高效的文档转换工具，将 Word 文档转换为结构化的 YAML 格式。')
    
    # Add a list
    doc.add_paragraph('项目特点：', style='Heading 2')
    doc.add_paragraph('• 支持多种文档元素', style='List Bullet')
    doc.add_paragraph('• 保持原始结构', style='List Bullet')
    doc.add_paragraph('• AI 友好的输出格式', style='List Bullet')
    
    # Add a table
    doc.add_heading('团队成员', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Add table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '姓名'
    hdr_cells[1].text = '职责'
    hdr_cells[2].text = '备注'
    
    # Add table data
    row_cells = table.add_row().cells
    row_cells[0].text = '张三'
    row_cells[1].text = '产品经理'
    row_cells[2].text = '负责需求分析和产品设计'
    
    row_cells = table.add_row().cells
    row_cells[0].text = '李四'
    row_cells[1].text = '开发工程师'
    row_cells[2].text = '负责核心功能开发'
    
    row_cells = table.add_row().cells
    row_cells[0].text = '王五'
    row_cells[1].text = '测试工程师'
    row_cells[2].text = '负责质量保证'
    
    # Add another section
    doc.add_heading('技术栈', level=2)
    doc.add_paragraph('后端技术：', style='Heading 3')
    doc.add_paragraph('• Python 3.8+', style='List Bullet')
    doc.add_paragraph('• python-docx 库', style='List Bullet')
    doc.add_paragraph('• PyYAML 库', style='List Bullet')
    
    doc.add_paragraph('AI 集成：', style='Heading 3')
    doc.add_paragraph('• OpenAI GPT-4o', style='List Bullet')
    doc.add_paragraph('• Langchain 框架', style='List Bullet')
    
    # Add final paragraph
    doc.add_paragraph('以上示例展示了 docx-plainify 工具能够处理的各种文档元素。')
    
    # Save the document
    doc.save('sample_document.docx')
    print("✅ 示例文档已创建: sample_document.docx")

if __name__ == '__main__':
    create_sample_docx() 